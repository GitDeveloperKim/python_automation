import docx

def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

#example 1
doc = docx.Document('demo.docx') #document 객체
len(doc.paragraphs) # paragraph 객체리스트를 가지고 있다 (문단)
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[1].runs)
print(doc.paragraphs[1].runs[0].text) # runs 객체의 리스트
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(getText('data/demo.docx'))

#style -> deprecated (https://m.blog.naver.com/sangja84/221530974938)
#print(doc.paragraphs[0].style)
#doc.paragraphs[0].style = 'Normal'
#doc.paragraphs[1].runs[0].style = 'QuoteChar'
#doc.paragraphs[1].runs[1].underline = True
#doc.paragraphs[1].runs[3].underline = True
#doc.save('restyled.docx')

# save
doc = docx.Document()
doc.add_paragraph('hello world')
paraObject1 = doc.add_paragraph('this is a second paragraph.')
paraObject2 = doc.add_paragraph('this is a yet another paragraph')
paraObject1.add_run('\nthis text is being added to the second paragraph')
doc.save('helloworld.docx')